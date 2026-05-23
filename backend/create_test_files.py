from reportlab.pdfgen import canvas
from docx import Document

# Create dummy JD
c = canvas.Canvas("test_jd.pdf")
c.drawString(100, 750, "Job Title: Python Backend Developer")
c.drawString(100, 730, "Requirements:")
c.drawString(100, 710, "- 3+ years Python experience")
c.drawString(100, 690, "- FastAPI or Django REST framework")
c.drawString(100, 670, "- Experience with PostgreSQL")
c.drawString(100, 650, "- REST API design")
c.drawString(100, 630, "- Git version control")
c.save()

# Create dummy resume 1
doc1 = Document()
doc1.add_heading("John Doe", 0)
doc1.add_paragraph("Email: john@email.com | Python Developer")
doc1.add_heading("Experience", level=1)
doc1.add_paragraph("4 years Python, FastAPI, PostgreSQL, Docker, REST APIs")
doc1.add_heading("Education", level=1)
doc1.add_paragraph("BSc Computer Science, 2020")
doc1.save("resume_john.docx")

# Create dummy resume 2
doc2 = Document()
doc2.add_heading("Jane Smith", 0)
doc2.add_paragraph("Email: jane@email.com | Full Stack Developer")
doc2.add_heading("Experience", level=1)
doc2.add_paragraph("2 years JavaScript, React, Node.js, some Python scripting")
doc2.add_heading("Education", level=1)
doc2.add_paragraph("BSc Information Technology, 2022")
doc2.save("resume_jane.docx")

print("Test files created successfully")