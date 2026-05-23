from reportlab.pdfgen import canvas
from docx import Document
import os

os.makedirs("demo_files", exist_ok=True)

def make_pdf(filename, lines):
    c = canvas.Canvas(f"demo_files/{filename}")
    y = 750
    for line in lines:
        c.drawString(50, y, line)
        y -= 20
    c.save()

def make_docx(filename, sections):
    doc = Document()
    for heading, content in sections:
        if heading:
            doc.add_heading(heading, level=1 if heading.isupper() else 2)
        if content:
            doc.add_paragraph(content)
    doc.save(f"demo_files/{filename}")

# Job Description
make_pdf("job_description.pdf", [
    "Position: Senior Python Backend Developer",
    "Company: TechCorp Malaysia Sdn Bhd",
    "",
    "About the Role:",
    "We are looking for a Senior Python Developer to join our growing",
    "engineering team. You will build and maintain scalable backend services.",
    "",
    "Requirements:",
    "- 4+ years of Python development experience",
    "- Strong experience with FastAPI or Django REST Framework",
    "- PostgreSQL and database optimization",
    "- RESTful API design and development",
    "- Docker and containerization",
    "- Git version control",
    "- Experience with cloud platforms (AWS or GCP preferred)",
    "",
    "Nice to Have:",
    "- Redis for caching",
    "- Celery for task queues",
    "- CI/CD pipeline experience",
    "",
    "Salary: RM 8,000 - RM 12,000 per month",
    "Location: Kuala Lumpur (Hybrid)"
])

# Strong match - Senior dev
make_docx("ahmad_rizal.docx", [
    ("Ahmad Rizal bin Zainudin", None),
    (None, "Email: ahmad.rizal@email.com | LinkedIn: linkedin.com/in/ahmadrizal"),
    (None, "Senior Python Developer | 6 Years Experience | Kuala Lumpur"),
    ("PROFESSIONAL SUMMARY", None),
    (None, "Senior Python developer with 6 years building scalable backend systems for fintech and e-commerce platforms. Strong advocate for clean code and API-first design."),
    ("TECHNICAL SKILLS", None),
    (None, "Languages: Python (6 years), SQL, Bash"),
    (None, "Frameworks: FastAPI, Django REST Framework, Flask"),
    (None, "Databases: PostgreSQL, MySQL, Redis"),
    (None, "DevOps: Docker, Kubernetes, AWS EC2, S3, GitHub Actions CI/CD"),
    (None, "Tools: Git, Celery, RabbitMQ, Nginx"),
    ("WORK EXPERIENCE", None),
    (None, "Senior Backend Developer - FinPay Malaysia (2021 - Present)"),
    (None, "Built FastAPI microservices handling 50,000 daily transactions. Optimized PostgreSQL queries reducing load time by 40%. Led team of 3 junior developers."),
    (None, "Backend Developer - ShopEasy (2019 - 2021)"),
    (None, "Developed Django REST APIs for e-commerce platform. Implemented Redis caching layer. Containerized services with Docker."),
    ("EDUCATION", None),
    (None, "BSc Computer Science - Universiti Malaya (2018)"),
])

# Good match - meets most requirements
make_docx("priya_nair.docx", [
    ("Priya Nair", None),
    (None, "Email: priya.nair@email.com | GitHub: github.com/priyanair"),
    (None, "Backend Developer | 4 Years Experience"),
    ("SUMMARY", None),
    (None, "Python developer with 4 years experience building REST APIs and backend systems. Comfortable with FastAPI and PostgreSQL. Looking to grow into senior roles."),
    ("SKILLS", None),
    (None, "Python (4 years), FastAPI, PostgreSQL, Git, Docker basics, REST APIs, Linux"),
    ("EXPERIENCE", None),
    (None, "Backend Developer - DataSync Sdn Bhd (2021 - Present)"),
    (None, "Built and maintained FastAPI services for data pipeline platform. Worked with PostgreSQL databases, wrote complex queries and migrations."),
    (None, "Junior Developer - Webcraft Agency (2020 - 2021)"),
    (None, "Built Django applications and REST endpoints for client projects."),
    ("EDUCATION", None),
    (None, "BSc Software Engineering - UTM (2020)"),
])

# Partial match - frontend dev trying backend
make_docx("wei_liang_tan.docx", [
    ("Tan Wei Liang", None),
    (None, "Email: weiliang@email.com"),
    (None, "Full Stack Developer | 3 Years Experience"),
    ("SUMMARY", None),
    (None, "Full stack developer with strong React and Node.js skills. Recently started learning Python and Django. Comfortable with REST APIs from the frontend perspective."),
    ("SKILLS", None),
    (None, "JavaScript (3 years), React, Node.js, Express, Python (6 months), Django (basic), MySQL, Git, HTML, CSS"),
    ("EXPERIENCE", None),
    (None, "Full Stack Developer - Creative Studio KL (2022 - Present)"),
    (None, "Built React frontends and Node.js backends for client websites. Integrated third-party REST APIs. Basic MySQL database management."),
    ("EDUCATION", None),
    (None, "Diploma in Information Technology - Sunway College (2021)"),
])

# Career changer - data scientist with some Python
make_docx("nurul_ain.docx", [
    ("Nurul Ain binti Hassan", None),
    (None, "Email: nurulain@email.com"),
    (None, "Data Scientist | 3 Years Experience"),
    ("SUMMARY", None),
    (None, "Data scientist with strong Python skills in data analysis and machine learning. No formal backend development experience but strong Python fundamentals and SQL knowledge."),
    ("SKILLS", None),
    (None, "Python (3 years - data focus), Pandas, NumPy, Scikit-learn, SQL, PostgreSQL, Jupyter, Git, some Flask for model serving"),
    ("EXPERIENCE", None),
    (None, "Data Scientist - Analytics Corp (2022 - Present)"),
    (None, "Built ML models for churn prediction. Wrote Python scripts for data pipelines. Used PostgreSQL for data storage and querying. Deployed one Flask API for model serving."),
    ("EDUCATION", None),
    (None, "BSc Statistics - UPM (2021)"),
])

# Fresh grad - strong academics, no experience
make_docx("kevin_lim.docx", [
    ("Kevin Lim Jun Kai", None),
    (None, "Email: kevinlim@email.com | GitHub: github.com/kevinlim"),
    (None, "Fresh Graduate | Python Enthusiast"),
    ("SUMMARY", None),
    (None, "Recent CS graduate with strong academic background in Python and software engineering. Built several personal projects using FastAPI and PostgreSQL. No professional experience yet."),
    ("SKILLS", None),
    (None, "Python, FastAPI, PostgreSQL, Git, Docker (basic), REST APIs, HTML, CSS"),
    ("PROJECTS", None),
    (None, "Student Management System - Built with FastAPI and PostgreSQL. Full CRUD operations, JWT authentication, deployed on Railway."),
    (None, "Task API - Simple REST API built with FastAPI. Documented with Swagger."),
    ("EDUCATION", None),
    (None, "BSc Computer Science - MMU Cyberjaya (2024) | CGPA 3.7"),
])

# Weak match - unrelated background
make_docx("siti_rohani.docx", [
    ("Siti Rohani binti Mahmud", None),
    (None, "Email: sitirohani@email.com"),
    (None, "IT Support Specialist | 5 Years Experience"),
    ("SUMMARY", None),
    (None, "Experienced IT support specialist with strong troubleshooting skills. Basic knowledge of Python scripting for automation tasks. Looking to transition into development."),
    ("SKILLS", None),
    (None, "Windows Server, Active Directory, Network troubleshooting, Python (basic scripting), SQL (basic queries), Excel, JIRA"),
    ("EXPERIENCE", None),
    (None, "IT Support Specialist - Government Agency (2019 - Present)"),
    (None, "Managed IT infrastructure for 200+ users. Wrote basic Python scripts for file automation. Handled helpdesk tickets and network issues."),
    ("EDUCATION", None),
    (None, "Diploma in Computer Science - Polytechnic (2018)"),
])

print("Demo files created in backend/demo_files/")
print("Files:")
print("  - job_description.pdf")
print("  - ahmad_rizal.docx      (Strong Hire - 6yr senior dev)")
print("  - priya_nair.docx       (Hire - meets requirements)")
print("  - wei_liang_tan.docx    (Consider - frontend trying backend)")
print("  - nurul_ain.docx        (Borderline - data scientist)")
print("  - kevin_lim.docx        (Consider - strong fresh grad)")
print("  - siti_rohani.docx      (Do Not Hire - IT support)")