from reportlab.pdfgen import canvas
from docx import Document
import os

os.makedirs("demo_files", exist_ok=True)

def make_pdf(filename, lines):
    c = canvas.Canvas(f"demo_files/{filename}")
    y = 750
    for line in lines:
        c.drawString(50, y, line)
        y -= 18
    c.save()

def make_docx(filename, sections):
    doc = Document()
    for heading, content in sections:
        if heading:
            doc.add_heading(heading, level=1 if heading.isupper() else 2)
        if content:
            doc.add_paragraph(content)
    doc.save(f"demo_files/{filename}")

# JD - tailored to match Krish's actual profile
make_pdf("job_description.pdf", [
    "Position: AI Engineer / LLM Application Developer",
    "Company: NovaMind AI Sdn Bhd, Kuala Lumpur",
    "",
    "About the Role:",
    "We are building next-generation AI-powered products and are looking for a talented",
    "AI Engineer to join our team. You will design and build agentic AI systems,",
    "integrate large language models into production pipelines, and develop scalable",
    "backend services that power intelligent automation workflows.",
    "",
    "Requirements:",
    "- Strong Python development skills",
    "- Experience building AI agents or agentic pipelines (LangGraph, AutoGen, or similar)",
    "- Hands-on experience with LLM integration (OpenAI API, Ollama, or similar)",
    "- RAG pipeline development (vector databases, embeddings, semantic search)",
    "- FastAPI or equivalent backend framework for REST API development",
    "- Experience with ML frameworks (PyTorch, TensorFlow, or scikit-learn)",
    "- Git version control and open-source project experience",
    "- Familiarity with prompt engineering techniques",
    "",
    "Nice to Have:",
    "- MLflow or equivalent experiment tracking",
    "- Electron or desktop application packaging",
    "- Linux system administration experience",
    "- Published VS Code extensions or open-source contributions",
    "- Research experience in AI or optimization",
    "",
    "We welcome strong fresh graduates and students with compelling project portfolios.",
    "Location: Kuala Lumpur | Mode: Hybrid | Available for internship consideration",
])

# Krish - Strong Hire (tailored to his actual CV)
make_docx("krish_sharma.docx", [
    ("Krish Sharma", None),
    (None, "Email: tp084234@mail.apu.edu.my | GitHub: github.com/krishsharma5169"),
    (None, "AI Engineer and Application Developer | BSc Computer Science (AI) | CGPA 3.82"),
    ("SUMMARY", None),
    (None, "CS/AI student at Asia Pacific University with strong hands-on experience building agentic AI systems, LLM integrations, and RAG pipelines. Projects span multi-agent orchestration, local LLM inference, neural networks, and FastAPI backend development. Active open-source contributor with published VS Code extension."),
    ("TECHNICAL PROJECTS", None),
    (None, "FinSight: Agentic Financial Research System (Apr 2026 - May 2026)"),
    (None, "Built a 5-agent LangGraph pipeline with conditional state graph routing and circuit breaker pattern. Integrated FinBERT sentiment analysis over live financial news. Tracked model iterations with MLflow. Built interactive Streamlit dashboard with Plotly. Tech: Python, LangGraph, GPT-4o-mini, FinBERT, MLflow, OpenAI API."),
    (None, "BonBon AI Coding Assistant (Jan 2026 - Mar 2026)"),
    (None, "Built fully local AI coding assistant with custom agentic pipeline featuring repair loop and critic agent. Implemented RAG using ChromaDB and nomic-embed-text embeddings, reducing solve time from 117s to 28s. FastAPI backend with intelligent mode detection routing. Packaged as Electron desktop app. Tech: Python, FastAPI, Ollama, Qwen2.5-Coder, ChromaDB."),
    (None, "VisionDebug: AI-Powered Python Debugger for VS Code (Jan 2026 - Mar 2026)"),
    (None, "Built AI debugging system combining AST-based static analysis and local LLM inference. FastAPI backend with AST analysis and error classification pipelines. Published as open-source VS Code extension on the Marketplace. Tech: Python, FastAPI, Uvicorn, Ollama, DeepSeek Coder."),
    (None, "Network Intrusion Detection Neural Network (Mar 2026)"),
    (None, "Built PyTorch feedforward neural network on UNSW-NB15 dataset achieving 99.5% accuracy and 0.980 macro F1. Handled severe class imbalance using BCELoss pos_weight. Tech: Python, PyTorch, scikit-learn, pandas."),
    ("SKILLS", None),
    (None, "AI/LLM: LangGraph, Agentic AI, RAG Pipelines, OpenAI API, Ollama, ChromaDB, Vector Databases, Prompt Engineering, FinBERT"),
    (None, "ML: PyTorch, TensorFlow, Keras, scikit-learn, MLflow, Neural Networks, Feature Engineering"),
    (None, "Backend: Python, FastAPI, REST APIs, Uvicorn, Git, GitHub"),
    (None, "Systems: Linux (Ubuntu, Rocky Linux), MySQL, Microsoft SQL Server"),
    (None, "Frontend: Electron, Streamlit, JavaScript, HTML, CSS"),
    ("EDUCATION", None),
    (None, "BSc (Hons) Computer Science with Artificial Intelligence - Asia Pacific University (Expected Nov 2027) | CGPA 3.82"),
    ("CERTIFICATIONS", None),
    (None, "Red Hat System Administration II (RH134), Red Hat, 2026"),
    (None, "Red Hat System Administration I (RH124), Red Hat, 2025"),
    (None, "Plug-and-Play AI: How to Build Applications with Chutes.ai, APUAIC, 2025"),
])

# Strong candidate 2 - Senior AI Engineer, strong match
make_docx("aisha_rahman.docx", [
    ("Aisha Rahman", None),
    (None, "Email: aisha.rahman@email.com | GitHub: github.com/aisharahman"),
    (None, "Senior AI Engineer | 5 Years Experience | Kuala Lumpur"),
    ("SUMMARY", None),
    (None, "Senior AI engineer with 5 years building production LLM applications and agentic systems. Strong background in RAG pipelines, OpenAI API integration, and FastAPI backend services. Led AI product teams at two Malaysian tech companies."),
    ("SKILLS", None),
    (None, "Python (5 years), LangChain, LangGraph, OpenAI API, RAG Pipelines, ChromaDB, Pinecone, FastAPI, PyTorch, MLflow, Docker, AWS, Git"),
    ("EXPERIENCE", None),
    (None, "Senior AI Engineer - TechNova Malaysia (2022 - Present)"),
    (None, "Built multi-agent LLM pipelines for enterprise document processing. Designed RAG systems with vector search over 100k+ document corpora. Led team of 3 engineers. Deployed FastAPI services on AWS."),
    (None, "AI Engineer - DataMind Sdn Bhd (2020 - 2022)"),
    (None, "Integrated OpenAI API into SaaS product. Built prompt engineering frameworks. Developed Python ML pipelines with MLflow experiment tracking."),
    ("EDUCATION", None),
    (None, "BSc Computer Science - Universiti Malaya (2019) | CGPA 3.75"),
])

# Good match - ML engineer, partial fit
make_docx("rajan_pillai.docx", [
    ("Rajan Pillai", None),
    (None, "Email: rajan.pillai@email.com"),
    (None, "Machine Learning Engineer | 3 Years Experience"),
    ("SUMMARY", None),
    (None, "ML engineer with 3 years in model training and deployment. Strong PyTorch and scikit-learn background. Some LLM integration experience via OpenAI API but limited agentic pipeline development. No RAG or vector database experience."),
    ("SKILLS", None),
    (None, "Python (3 years), PyTorch, TensorFlow, scikit-learn, OpenAI API (basic), FastAPI (basic), SQL, Git, Docker"),
    ("EXPERIENCE", None),
    (None, "ML Engineer - Analytics Firm KL (2022 - Present)"),
    (None, "Trained and deployed classification and regression models. Built Python data pipelines. Basic FastAPI deployment for model serving. Some OpenAI API experimentation."),
    ("EDUCATION", None),
    (None, "BSc Data Science - UTM (2021)"),
])

# Partial match - backend dev, no AI experience
make_docx("mei_ling_wong.docx", [
    ("Wong Mei Ling", None),
    (None, "Email: meiling.wong@email.com"),
    (None, "Backend Developer | 4 Years Experience"),
    ("SUMMARY", None),
    (None, "Experienced Python backend developer with strong FastAPI and REST API skills. No AI, LLM, or machine learning experience. Interested in transitioning into AI engineering but no hands-on background yet."),
    ("SKILLS", None),
    (None, "Python (4 years), FastAPI, Django, PostgreSQL, Redis, Docker, Git, REST APIs, Linux"),
    ("EXPERIENCE", None),
    (None, "Backend Developer - Webservices Corp (2021 - Present)"),
    (None, "Built and maintained FastAPI microservices. PostgreSQL database design and optimization. Docker containerization. REST API design and documentation."),
    ("EDUCATION", None),
    (None, "BSc Software Engineering - MMU (2020)"),
])

# Weak match - data analyst, surface-level Python
make_docx("hafiz_nordin.docx", [
    ("Hafiz bin Nordin", None),
    (None, "Email: hafiz.nordin@email.com"),
    (None, "Data Analyst | 3 Years Experience"),
    ("SUMMARY", None),
    (None, "Data analyst with Python and SQL skills focused on reporting and dashboards. No LLM, agentic AI, RAG, or backend API development experience. Uses Python primarily for data analysis with Pandas and Excel automation."),
    ("SKILLS", None),
    (None, "Python (Pandas, NumPy, basic), SQL, Excel, Power BI, Tableau, Git (basic)"),
    ("EXPERIENCE", None),
    (None, "Data Analyst - Finance Company KL (2022 - Present)"),
    (None, "Built dashboards and reports using Power BI and Tableau. Python scripts for data cleaning and Excel automation. SQL queries for data extraction."),
    ("EDUCATION", None),
    (None, "BSc Business Analytics - Sunway University (2021)"),
])

# No match - IT support
make_docx("faridah_hassan.docx", [
    ("Faridah binti Hassan", None),
    (None, "Email: faridah.hassan@email.com"),
    (None, "IT Support Specialist | 6 Years Experience"),
    ("SUMMARY", None),
    (None, "Experienced IT support specialist managing infrastructure and helpdesk operations. No programming, AI, or software development experience. Looking to explore tech roles."),
    ("SKILLS", None),
    (None, "Windows Server, Active Directory, Network Troubleshooting, JIRA, Excel, basic Python scripting for file tasks"),
    ("EXPERIENCE", None),
    (None, "IT Support Lead - Government Agency (2019 - Present)"),
    (None, "Managed IT infrastructure for 300+ users. Handled helpdesk tickets. Basic Python scripts for file renaming automation. No software development or AI experience."),
    ("EDUCATION", None),
    (None, "Diploma in Information Technology - Polytechnic Shah Alam (2018)"),
])

print("Demo files created in demo_files/")
print("")
print("Files:")
print("  job_description.pdf    - AI Engineer / LLM Application Developer role")
print("  krish_sharma.docx      - Strong Hire (your actual profile)")
print("  aisha_rahman.docx      - Strong Hire (senior AI engineer)")
print("  rajan_pillai.docx      - Consider (ML engineer, partial fit)")
print("  mei_ling_wong.docx     - Borderline (backend dev, no AI)")
print("  hafiz_nordin.docx      - Do Not Hire (data analyst)")
print("  faridah_hassan.docx    - Do Not Hire (IT support)")